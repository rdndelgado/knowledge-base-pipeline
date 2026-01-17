import os
import pandas as pd
from datetime import datetime
from docx import Document
from utils.logger import logger
from utils.tokens import count_tokens
from typing import List, Dict, Tuple
import nltk
from nltk.tokenize import sent_tokenize
from models.database import KBDocument

# Download NLTK data if not already present
try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    logger.info("Downloading NLTK punkt_tab tokenizer...")
    nltk.download('punkt_tab', quiet=True)

class DocumentService:
    
    def __init__(self, docs_dir="documents", error_log_path="logs/errors.csv", mysql_service=None):
        self.docs_dir = docs_dir
        self.error_log_path = error_log_path
        self.mysql_service = mysql_service
        os.makedirs(os.path.dirname(error_log_path), exist_ok=True)

        # Chunking parameters
        self.CHUNK_MIN_SIZE = 500
        self.CHUNK_MAX_SIZE = 800
        self.CHUNK_OVERLAP = 80
    
    def process(self, documents: List[str] = None) -> Tuple[List[Dict], List[Dict]]:
        """
            Map .docx files to existing database records and update their content.
            Generate chunks for the documents.
            Returns a list of records that were successfully mapped with chunks.
        """
        mapped_docs = self._map_documents(documents=documents)
        chunked_docs = self._generate_chunks(mapped_docs)
        
        return mapped_docs, chunked_docs

    def _map_documents(self, documents: List[str] = None) -> List[Dict]:
        """
        Map .docx files to existing database records and update their content.
        Returns a list of records that were successfully updated.
        
        Parameters:
            documents (List[str], optional): List of filenames to process.
                If None, all .docx files in the directory are considered.
        """
        
        total_docs = 0
        mapped_documents = []

        # List all .docx files in the directory
        all_files = [f for f in os.listdir(self.docs_dir) if f.endswith(".docx")]

        # Filter files if documents parameter is provided
        files_to_process = [f for f in all_files if not documents or f in documents]

        if not files_to_process:
            logger.warning("No documents to process.")
            return []

        for filename in files_to_process:
            total_docs += 1
            title = os.path.splitext(filename)[0]
            file_path = os.path.join(self.docs_dir, filename)

            # Read .docx content
            content = self._read_docx(file_path)
            if not content.strip():
                logger.warning(f"Empty content for: {title}")
                continue

            try:
                # Fetch existing document from the database
                document = self.mysql_service.get_document_by_title(title)

                # Read .docx content
                content = self._read_docx(file_path)
                if not content.strip():
                    logger.warning(f"Empty content for: {title}")
                    continue
                
                if not document:
                    logger.warning(f"Document not found in the database: {title}")
                    payload = KBDocument(
                        title=title,
                        content=content
                    )

                    # if document is new, insert one.
                    data = self.mysql_service.insert_document(payload)
                    if data:
                        logger.info(f"Inserted new document: {title}")
                        mapped_documents.append(data)
                    else:
                        logger.warning(f"Failed to insert new document: {title}")
                    continue # Proceed or Skip to next document

                # Update the content in the document
                document["content"] = content
                mapped_documents.append(document)
            except Exception as e:
                self._log_error(title, str(e))
                logger.error(f"Error while mapping document {title}.")
                raise
        
        logger.info(f"Mapped {len(mapped_documents)}/{total_docs} documents.")
        return mapped_documents
    
    def _generate_chunks(self, documents: List[Dict]) -> List[Dict]:
        """
        Generate chunks for each document based on chunk size.
        Returns a flat list of chunks with document_id included.
        """
        all_chunks = []

        for _, doc in enumerate(documents):
            doc_id = doc["id"]
            content = str(doc["content"])
            title = doc.get("title")

            chunks = self._chunk_text(content)
            for idx, chunk in enumerate(chunks):
                token_count = count_tokens(chunk)
                all_chunks.append({
                    "document_id": doc_id,
                    "chunk_index": idx,
                    "content": chunk,
                    "token_count": token_count
                })
            
            logger.info(f"Generated {len(chunks)} chunks for document: {title}")
        
        logger.info(f"Total generated {len(all_chunks)} chunks across all documents")
        return all_chunks

    def _chunk_text(self, text: str):
        """
        Split text into chunks with overlap and sentence boundaries.
        """
        sentences = sent_tokenize(text)
        chunks = []
        current_chunk = []
        current_tokens = 0

        for sentence in sentences:
            sentence_tokens = count_tokens(sentence)

            # If adding this sentence would exceed max tokens → start new chunk
            if current_tokens + sentence_tokens > self.CHUNK_MAX_SIZE:
                if current_chunk:
                    chunks.append(" ".join(current_chunk))
                
                # Start overlap
                overlap_tokens = 0
                overlap_chunk = []
                while current_chunk and overlap_tokens < self.CHUNK_OVERLAP:
                    sent = current_chunk.pop()
                    overlap_chunk.insert(0, sent)
                    overlap_tokens += count_tokens(sent)

                current_chunk = overlap_chunk + [sentence]
                current_tokens = overlap_tokens + sentence_tokens
            else:
                current_chunk.append(sentence)
                current_tokens += sentence_tokens

        # Add last chunk
        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks
    
    def _read_docx(self, file_path):
        """Read .docx file and return plain text."""
        try:
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.error(f"Failed to read {file_path}.")
            self._log_error(file_path, str(e))
            raise

    def _log_error(self, filename, error_message):
        """Append error details to error log CSV with timestamp."""
        try:
            timestamp = datetime.now().isoformat()
            df = pd.DataFrame([
                {
                    "timestamp": timestamp,
                    "filename": filename,
                    "error": error_message
                }
            ])

            if os.path.exists(self.error_log_path):
                df.to_csv(self.error_log_path, mode="a", index=False, header=False)
            else:
                df.to_csv(self.error_log_path, index=False)
        except Exception as e:
            logger.error(f"Failed to write to error log: {e}")
            raise